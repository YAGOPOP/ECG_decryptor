from docx import Document
from docx.shared import Inches
from datetime import date

# Collecting data
name = input("Ф.И.О.: ")
bdr = input("Дата рождения (ДД.ММ.ГГГГ): ")
pulse = input("ЧСС: ")
alpha = input("α: ")
P = input("P: ")
PQ = input("PQ: ")
QRS = input("QTS: ")
QT = input("QT: ")
QTc = input("QTs: ")

# Processing data
todate = date.today()
bd = list(map(int, bdr.split(".")))
birthday = date(bd[2], bd[1], bd[0])
age = int((todate - birthday).days / 365.25)






replacements = {
    "DD.MM.YYYY": f"{str(todate.day).zfill(2)}.{str(todate.month).zfill(2)}.{todate.year}",
    "Ф.И.О.:     ": f"Ф.И.О.: {name}",
    "Возраст:    ": f"Возраст: {bdr} / {age} лет",
    "ЧСС уд/мин": f"ЧСС {pulse} уд/мин",
    "α=   °": f"α= {alpha}°",
    "P - 0,  с": "P - 0,12 с",
    "PQ - 0,  с": "PQ - 0,18 с",
    "QRS -  0,  с": "QRS - 0,08 с",
    "QT -  0,  с": "QT - 0,36 с",
    "QTc - 0,  с": "QTc - 0,40 с",
}




# Working with document
doc = Document("./templates/экг домодедово.docx")

for par in doc.paragraphs:
    input(par.text)

# for par in doc.paragraphs:
#     for k, v in replacements.items():
#         if par.text
#         print([r.text for r in par.runs])

# doc.save("./res.docx")